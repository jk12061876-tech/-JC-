# -*- coding: utf-8 -*-
"""
Product Catalog Generator
Burgundy Premium Color Scheme
A3 Landscape with 9-column product tables
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Premium Color Scheme
CREAM = "EFDFCE"          # Cream - Background
BURGUNDY = "7E041D"       # Burgundy Red - Accent
DARK_GRAY = "363636"      # Dark Gray - Text
WHITE = "FFFFFF"

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, color="7E041D", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="nil"/>'
        f'<w:left w:val="nil"/>'
        f'<w:bottom w:val="nil"/>'
        f'<w:right w:val="nil"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def create_cover_page(doc):
    """Create cover page with cream background"""
    bg_table = doc.add_table(rows=1, cols=1)
    bg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bg_cell = bg_table.cell(0, 0)
    set_cell_shading(bg_cell, CREAM)
    set_no_border(bg_cell)

    para = bg_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para.add_run("\n\n\n\n\n")

    logo_run = para.add_run("[ LOGO ]\n\n\n")
    logo_run.font.size = Pt(24)
    logo_run.font.color.rgb = RGBColor(126, 4, 29)
    logo_run.font.name = 'Arial'

    company_run = para.add_run("YOUR COMPANY NAME\n")
    company_run.font.size = Pt(48)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(54, 54, 54)
    company_run.font.name = 'Arial'

    line_run = para.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n")
    line_run.font.size = Pt(14)
    line_run.font.color.rgb = RGBColor(126, 4, 29)

    title_run = para.add_run("PRODUCT CATALOG\n")
    title_run.font.size = Pt(60)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(126, 4, 29)
    title_run.font.name = 'Arial'

    sub_run = para.add_run("Server Components & Solutions\n\n")
    sub_run.font.size = Pt(24)
    sub_run.font.color.rgb = RGBColor(54, 54, 54)
    sub_run.font.name = 'Arial'

    year_run = para.add_run("2024 - 2025\n\n\n\n\n")
    year_run.font.size = Pt(20)
    year_run.font.color.rgb = RGBColor(126, 4, 29)

    bottom_line = para.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n")
    bottom_line.font.size = Pt(10)
    bottom_line.font.color.rgb = RGBColor(126, 4, 29)

    contact_run = para.add_run("Tel: XXX-XXXX-XXXX  |  Email: example@company.com\n")
    contact_run.font.size = Pt(14)
    contact_run.font.color.rgb = RGBColor(54, 54, 54)
    contact_run.font.name = 'Arial'

    address_run = para.add_run("Address: Your Company Address\n\n\n")
    address_run.font.size = Pt(14)
    address_run.font.color.rgb = RGBColor(54, 54, 54)
    address_run.font.name = 'Arial'

    doc.add_page_break()

def create_company_intro_page(doc):
    """Create company introduction page"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ABOUT US")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    doc.add_paragraph()

    intro_table = doc.add_table(rows=1, cols=1)
    intro_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = intro_table.cell(0, 0)
    set_cell_shading(cell, CREAM)
    set_cell_border(cell, BURGUNDY, "8")

    cell_para = cell.paragraphs[0]
    run = cell_para.add_run("""
    Company Introduction

    [Enter your company introduction here]

    • Year of establishment and development history
    • Main products and service scope
    • Company size and team strength
    • Core competitive advantages
    • Service philosophy

    """)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(54, 54, 54)
    run.font.name = 'Arial'

    doc.add_paragraph()

    adv_title = doc.add_paragraph()
    adv_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = adv_title.add_run("OUR ADVANTAGES")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    doc.add_paragraph()

    adv_table = doc.add_table(rows=2, cols=2)
    adv_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    advantages = [
        ("Genuine Products", "Original products with full warranty"),
        ("Competitive Pricing", "Direct from manufacturer"),
        ("Technical Support", "Professional pre & after sales service"),
        ("Fast Delivery", "Same-day shipping available")
    ]

    for i, (title_text, desc) in enumerate(advantages):
        row = i // 2
        col = i % 2
        cell = adv_table.cell(row, col)
        set_cell_shading(cell, CREAM)
        set_cell_border(cell, DARK_GRAY, "4")
        para = cell.paragraphs[0]

        title_run = para.add_run("■ " + title_text + "\n")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(126, 4, 29)
        title_run.font.name = 'Arial'

        desc_run = para.add_run(desc)
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = RGBColor(54, 54, 54)
        desc_run.font.name = 'Arial'

    doc.add_page_break()

def create_product_section_header(doc, title, subtitle):
    """Create product section header"""
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    set_cell_shading(cell, BURGUNDY)
    set_no_border(cell)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run("  " + title + "  ")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Arial'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(54, 54, 54)
    run.font.name = 'Arial'

    doc.add_paragraph()

def create_product_table_9col(doc, products):
    """Create product table with 9 columns for memory products
    Columns: Image, Option Part Number, Capacity, Rank, Speed, CAS Latency, Technology, Warranty, Notes
    """
    columns = ["Image", "Option Part Number", "Capacity", "Rank", "Speed", "CAS Latency", "Technology", "Warranty", "Notes"]

    table = doc.add_table(rows=1, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, col_name in enumerate(columns):
        cell = header_row.cells[i]
        set_cell_shading(cell, BURGUNDY)
        set_cell_border(cell, BURGUNDY, "2")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(col_name)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Arial'

    # Data rows
    for idx, product in enumerate(products):
        row = table.add_row()
        bg_color = WHITE if idx % 2 == 0 else CREAM

        # Image placeholder cell (first column)
        img_cell = row.cells[0]
        set_cell_shading(img_cell, bg_color)
        set_cell_border(img_cell, BURGUNDY, "2")
        img_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        img_para = img_cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_para.add_run("[IMG]")
        img_run.font.size = Pt(7)
        img_run.font.color.rgb = RGBColor(126, 4, 29)
        img_run.font.name = 'Arial'

        # Product data columns (8 columns)
        for i, value in enumerate(product):
            cell = row.cells[i + 1]
            set_cell_shading(cell, bg_color)
            set_cell_border(cell, BURGUNDY, "2")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            run = para.add_run(str(value))
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(54, 54, 54)
            run.font.name = 'Arial'

    doc.add_paragraph()

def create_hdd_table_9col(doc, products):
    """Create product table with 9 columns for HDD products
    Columns: Image, Option Part Number, Capacity, Form Factor, Speed, Interface, Technology, Warranty, Notes
    """
    columns = ["Image", "Option Part Number", "Capacity", "Form Factor", "Speed", "Interface", "Technology", "Warranty", "Notes"]

    table = doc.add_table(rows=1, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, col_name in enumerate(columns):
        cell = header_row.cells[i]
        set_cell_shading(cell, BURGUNDY)
        set_cell_border(cell, BURGUNDY, "2")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(col_name)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Arial'

    # Data rows
    for idx, product in enumerate(products):
        row = table.add_row()
        bg_color = WHITE if idx % 2 == 0 else CREAM

        # Image placeholder cell
        img_cell = row.cells[0]
        set_cell_shading(img_cell, bg_color)
        set_cell_border(img_cell, BURGUNDY, "2")
        img_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        img_para = img_cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_para.add_run("[IMG]")
        img_run.font.size = Pt(7)
        img_run.font.color.rgb = RGBColor(126, 4, 29)
        img_run.font.name = 'Arial'

        # Product data columns
        for i, value in enumerate(product):
            cell = row.cells[i + 1]
            set_cell_shading(cell, bg_color)
            set_cell_border(cell, BURGUNDY, "2")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            run = para.add_run(str(value))
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(54, 54, 54)
            run.font.name = 'Arial'

    doc.add_paragraph()

def main():
    doc = Document()

    # Set A3 Landscape
    section = doc.sections[0]
    section.page_width = Mm(420)   # A3 width
    section.page_height = Mm(297)  # A3 height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Cover
    create_cover_page(doc)

    # Company intro
    create_company_intro_page(doc)

    # ==================== HPE Server Memory ====================
    create_product_section_header(doc, "HPE SERVER MEMORY", "DDR4/DDR5 Series")

    g10_title = doc.add_paragraph()
    run = g10_title.add_run("▶ G10 Series Memory")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    # HPE G10 products: Part No, Capacity, Rank, Speed, CAS Latency, Technology, Warranty, Notes
    hpe_g10_products = [
        ("P00924-B21", "32GB", "Dual Rank x4", "DDR4-2933", "CAS-21-21-21", "Registered", "3-Year", ""),
        ("815100-B21", "32GB", "Dual Rank x4", "DDR4-2666", "CAS-19-19-19", "Registered", "3-Year", ""),
        ("P00930-B21", "64GB", "Dual Rank x4", "DDR4-2933", "CAS-21-21-21", "Registered", "3-Year", ""),
        ("P00926-B21", "64GB", "Quad Rank x4", "DDR4-2933", "CAS-21-21-21", "Load Reduced", "3-Year", ""),
        ("815101-B21", "64GB", "Quad Rank x4", "DDR4-2666", "CAS-19-19-19", "Load Reduced", "3-Year", ""),
        ("P11040-B21", "128GB", "Quad Rank x4", "DDR4-2933", "CAS-24-21-21", "Load Reduced", "3-Year", ""),
        ("P00928-B21", "128GB", "Octal Rank x4", "DDR4-2933", "CAS-24-21-21", "3DS Load Reduced", "3-Year", ""),
        ("815102-B21", "128GB", "Octal Rank x4", "DDR4-2666", "CAS-22-19-19", "3DS Load Reduced", "3-Year", ""),
    ]
    create_product_table_9col(doc, hpe_g10_products)

    g10plus_title = doc.add_paragraph()
    run = g10plus_title.add_run("▶ G10+ Series Memory")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    hpe_g10plus_products = [
        ("P06033-B21", "32GB", "Dual Rank x4", "DDR4-3200", "CAS-22-22-22", "Registered", "3-Year", ""),
        ("P06035-B21", "64GB", "Dual Rank x4", "DDR4-3200", "CAS-22-22-22", "Registered", "3-Year", ""),
        ("P40007-B21", "32GB", "Single Rank x4", "DDR4-3200", "CAS-22-22-22", "Registered", "3-Year", ""),
        ("P06037-B21", "128GB", "Quad Rank x4", "DDR4-3200", "CAS-22-22-22", "Load Reduced", "3-Year", ""),
        ("P06039-B21", "256GB", "Octal Rank x4", "DDR4-3200", "CAS-26-22-22", "3DS Load Reduced", "3-Year", ""),
    ]
    create_product_table_9col(doc, hpe_g10plus_products)

    g11_title = doc.add_paragraph()
    run = g11_title.add_run("▶ G11 Series Memory (DDR5)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    hpe_g11_products = [
        ("P64705-B21", "16GB", "Single Rank x8", "DDR5-5600", "CAS-46-45-45", "EC8 Registered", "3-Year", ""),
        ("P64706-B21", "32GB", "Dual Rank x8", "DDR5-5600", "CAS-46-45-45", "EC8 Registered", "3-Year", ""),
        ("P64707-B21", "64GB", "Dual Rank x4", "DDR5-5600", "CAS-46-45-45", "EC8 Registered", "3-Year", ""),
        ("P64708-B21", "96GB", "Dual Rank x4", "DDR5-5600", "CAS-45-45-45", "EC8 Registered", "3-Year", ""),
        ("P64709-B21", "128GB", "Quad Rank x4", "DDR5-5600", "CAS-52-45-45", "EC8 Registered 3DS", "3-Year", ""),
        ("P43322-B21", "16GB", "Single Rank x8", "DDR5-4800", "CAS-40-39-39", "EC8 Registered", "3-Year", ""),
        ("P43328-B21", "32GB", "Dual Rank x8", "DDR5-4800", "CAS-40-39-39", "EC8 Registered", "3-Year", ""),
        ("P43331-B21", "64GB", "Dual Rank x4", "DDR5-4800", "CAS-40-39-39", "EC8 Registered", "3-Year", ""),
        ("P66675-B21", "96GB", "2Rx4", "DDR5-4800", "-", "PC5-4800B-R", "3-Year", ""),
        ("P43334-B21", "128GB", "Quad Rank x4", "DDR5-4800", "CAS-46-39-39", "EC8 Registered 3DS", "3-Year", ""),
        ("P43337-B21", "256GB", "Octal Rank x4", "DDR5-4800", "CAS-46-39-39", "EC8 Registered 3DS", "3-Year", ""),
    ]
    create_product_table_9col(doc, hpe_g11_products)

    doc.add_page_break()

    # ==================== Dell Server Memory ====================
    create_product_section_header(doc, "DELL SERVER MEMORY", "DDR4/DDR5 Series")

    dell_products = [
        ("AA601616", "32GB", "2Rx4", "DDR4-2933", "-", "RDIMM", "3-Year", "R640/R740/R840/R940/T640"),
        ("AA601615", "64GB", "2Rx4", "DDR4-2933", "-", "RDIMM", "3-Year", "R640/R740/R840/R940/T640"),
        ("AA579531", "32GB", "2Rx4", "DDR4-2933", "-", "PC4-23400", "3-Year", "R640/R740/R840/R940/T640"),
        ("AA579530", "64GB", "2Rx4", "DDR4-2933", "-", "RDIMM", "3-Year", "R640/R740/R840/R940/T640"),
        ("AA783422", "32GB", "2Rx4", "DDR4-3200", "-", "RDIMM", "3-Year", "R6515/R6525/R740/R840/R940"),
        ("AA783423", "64GB", "2Rx4", "DDR4-3200", "-", "RDIMM", "3-Year", "R6515/R6525/R7515/C6525/R7525"),
        ("AA799110", "64GB", "2Rx4", "DDR4-3200", "-", "RDIMM", "3-Year", "R6515/R6525/R7515/C6525/R7525"),
        ("AA810828", "64GB", "2Rx4", "DDR4-3200", "-", "RDIMM", "3-Year", "R640/R740/R840/R940/T640"),
        ("AB445285", "128GB", "4Rx4", "DDR4-3200", "-", "LRDIMM", "3-Year", "R640/R740/R840/R940/T640"),
        ("AC239377", "16GB", "1Rx8", "DDR5-4800", "-", "RDIMM", "3-Year", "R660/R760/R6615/R7625"),
        ("AC239378", "32GB", "2Rx8", "DDR5-4800", "-", "RDIMM", "3-Year", "R660/R760/R6615/R7625"),
        ("AC239379", "64GB", "2Rx4", "DDR5-4800", "-", "RDIMM", "3-Year", "R660/R760/R860/R960"),
        ("AC830716", "16GB", "1Rx8", "DDR5-5600", "-", "RDIMM", "3-Year", ""),
        ("AC958788", "16GB", "1Rx8", "DDR5-5600", "-", "UDIMM", "3-Year", ""),
        ("AC774043", "32GB", "2Rx8", "DDR5-5600", "-", "ECC UDIMM", "3-Year", ""),
        ("AC830717", "32GB", "2Rx8", "DDR5-5600", "-", "ECC RDIMM", "3-Year", ""),
        ("AC830718", "64GB", "2Rx4", "DDR5-5600", "-", "RDIMM", "3-Year", ""),
        ("AC888060", "16GB", "-", "DDR5-5600", "-", "ECC RDIMM", "3-Year", ""),
    ]
    create_product_table_9col(doc, dell_products)

    doc.add_page_break()

    # ==================== Lenovo Server Memory ====================
    create_product_section_header(doc, "LENOVO SERVER MEMORY", "ThinkSystem Series")

    lenovo_products = [
        ("4ZC7A08707", "16GB", "1Rx4", "DDR4-2933", "-", "PC4-2933Y", "3-Year", "01KR353"),
        ("4ZC7A08709", "32GB", "2Rx4", "DDR4-2933", "-", "REG ECC", "3-Year", "01KR355"),
        ("4ZC7A08710", "64GB", "2Rx4", "DDR4-2933", "-", "RECC", "3-Year", "01KR356"),
        ("4X77A08634", "32GB", "2Rx8", "DDR4-3200", "-", "RDIMM", "3-Year", "02JK239"),
        ("4X77A08633", "32GB", "2Rx4", "DDR4-3200", "-", "PC4-3200AA", "3-Year", "02JK237"),
        ("4ZC7A15124", "64GB", "2Rx4", "DDR4-3200", "-", "RDIMM", "3-Year", "02JG340"),
        ("4X77A08635", "64GB", "2Vx4", "DDR4-3200", "-", "RDIMM", "3-Year", "02JK971"),
        ("4X77A77496", "32GB", "-", "DDR4-3200", "-", "ECC UDIMM", "3-Year", "03GX401"),
        ("4X77A85511", "16GB", "1Rx8", "DDR5-4800", "-", "TruDDR5", "3-Year", ""),
        ("4X77A77483", "32GB", "1Rx4", "DDR5-4800", "-", "TruDDR5", "3-Year", ""),
        ("4X77A88512", "32GB", "2Rx8", "DDR5-4800", "-", "TruDDR5", "3-Year", ""),
        ("4X77A77031", "32GB", "2Rx8", "DDR5-4800", "-", "TruDDR5", "3-Year", ""),
        ("4X77A81440", "32GB", "2Rx8", "DDR5-4800", "-", "TruDDR5", "3-Year", "03KL461"),
        ("4X77A81442", "64GB", "2Rx4", "DDR5-4800", "-", "TruDDR5", "3-Year", "03GX338"),
        ("4X77A77033", "64GB", "2Rx4", "DDR5-4800", "-", "TruDDR5", "3-Year", ""),
        ("4X77A77034", "128GB", "4Rx4", "DDR5-4800", "-", "3DS RDIMM", "3-Year", ""),
        ("4X77A77035", "256GB", "8Rx4", "DDR5-4800", "-", "TruDDR5", "3-Year", ""),
        ("4X77A88049", "32GB", "1Rx4", "DDR5-5600", "-", "TruDDR5", "3-Year", ""),
        ("4X77A88051", "32GB", "2Rx8", "DDR5-5600", "-", "TruDDR5", "3-Year", ""),
        ("4X77A90992", "64GB", "2Rx4", "DDR5-5600", "-", "TruDDR5", "3-Year", ""),
        ("4X77A88052", "64GB", "2Rx4", "DDR5-5600", "-", "TruDDR5", "3-Year", ""),
        ("4X77A93887", "128GB", "2Rx4", "DDR5-5600", "-", "RDIMM", "3-Year", ""),
        ("4X77A88054", "128GB", "4Rx4", "DDR5-5600", "-", "TruDDR5", "3-Year", ""),
        ("4X77A88055", "256GB", "8Rx4", "DDR5-5600", "-", "3DS RDIMM", "3-Year", ""),
    ]
    create_product_table_9col(doc, lenovo_products)

    doc.add_page_break()

    # ==================== HPE Enterprise HDD ====================
    create_product_section_header(doc, "HPE ENTERPRISE HDD", "SAS HDD Series")

    # Enterprise 10K SFF
    hdd_title1 = doc.add_paragraph()
    run = hdd_title1.add_run("▶ Enterprise 10K SFF (2.5-inch)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    # HDD: Part No, Capacity, Form Factor, Speed, Interface, Technology, Warranty, Notes
    hpe_10k_sff = [
        ("881457-B21", "2.4TB", "SFF 2.5\"", "10K", "SAS 12G", "512e DS", "3-Year", "SC"),
        ("872481-B21", "1.8TB", "SFF 2.5\"", "10K", "SAS 12G", "512e DS", "3-Year", "SC"),
        ("872479-B21", "1.2TB", "SFF 2.5\"", "10K", "SAS 12G", "DS firmware", "3-Year", "SC"),
        ("872477-B21", "600GB", "SFF 2.5\"", "10K", "SAS 12G", "DS firmware", "3-Year", "SC"),
        ("872475-B21", "300GB", "SFF 2.5\"", "10K", "SAS 12G", "DS firmware", "3-Year", "SC"),
    ]
    create_hdd_table_9col(doc, hpe_10k_sff)

    # Enterprise 15K SFF
    hdd_title2 = doc.add_paragraph()
    run = hdd_title2.add_run("▶ Enterprise 15K SFF (2.5-inch)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    hpe_15k_sff = [
        ("870759-B21", "900GB", "SFF 2.5\"", "15K", "SAS 12G", "DS firmware", "3-Year", "SC"),
        ("870757-B21", "600GB", "SFF 2.5\"", "15K", "SAS 12G", "DS firmware", "3-Year", "SC"),
        ("870753-B21", "300GB", "SFF 2.5\"", "15K", "SAS 12G", "DS firmware", "3-Year", "SC"),
    ]
    create_hdd_table_9col(doc, hpe_15k_sff)

    # Midline 7.2K SFF
    hdd_title3 = doc.add_paragraph()
    run = hdd_title3.add_run("▶ Midline 7.2K SFF (2.5-inch)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    hpe_72k_sff = [
        ("765466-B21", "2TB", "SFF 2.5\"", "7.2K", "SAS 12G", "512e", "1-Year", "SC"),
        ("832514-B21", "1TB", "SFF 2.5\"", "7.2K", "SAS 12G", "DS firmware", "1-Year", "SC"),
    ]
    create_hdd_table_9col(doc, hpe_72k_sff)

    # Midline 7.2K LFF
    hdd_title4 = doc.add_paragraph()
    run = hdd_title4.add_run("▶ Midline 7.2K LFF (3.5-inch)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    hpe_72k_lff = [
        ("P23863-B21", "16TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e ISE", "1-Year", "SC"),
        ("P09153-B21", "14TB", "LFF 3.5\"", "7.2K", "SAS 12G", "Helium 512e DS", "1-Year", "SC"),
        ("881779-B21", "12TB", "LFF 3.5\"", "7.2K", "SAS 12G", "Helium 512e DS", "1-Year", "SC"),
        ("857644-B21", "10TB", "LFF 3.5\"", "7.2K", "SAS 12G", "Helium 512e DS", "1-Year", "SC"),
        ("819201-B21", "8TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e DS", "1-Year", "SC"),
        ("861754-B21", "6TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e", "1-Year", "SC"),
    ]
    create_hdd_table_9col(doc, hpe_72k_lff)

    doc.add_page_break()

    # Mission Critical SFF BC
    hdd_title5 = doc.add_paragraph()
    run = hdd_title5.add_run("▶ Mission Critical 10K/15K SFF BC")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    hpe_mc_sff = [
        ("P28352-B21", "2.4TB", "SFF 2.5\"", "10K", "SAS 12G", "512e BC", "3-Year", "Mission Critical"),
        ("P28586-B21", "1.2TB", "SFF 2.5\"", "10K", "SAS 12G", "BC", "3-Year", "Mission Critical"),
        ("P28028-B21", "300GB", "SFF 2.5\"", "15K", "SAS 12G", "BC", "3-Year", "Mission Critical"),
        ("P40430-B21", "300GB", "SFF 2.5\"", "10K", "SAS 12G", "BC", "3-Year", "Mission Critical"),
    ]
    create_hdd_table_9col(doc, hpe_mc_sff)

    # Business Critical 7.2K LFF LP
    hdd_title6 = doc.add_paragraph()
    run = hdd_title6.add_run("▶ Business Critical 7.2K LFF LP (3.5-inch)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    hpe_bc_lff = [
        ("P37669-B21", "18TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e ISE", "1-Year", "LP"),
        ("P23608-B21", "16TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e ISE", "1-Year", "LP"),
        ("P09155-B21", "14TB", "LFF 3.5\"", "7.2K", "SAS 12G", "Helium 512e DS", "1-Year", "LP MDL"),
        ("881781-B21", "12TB", "LFF 3.5\"", "7.2K", "SAS 12G", "Helium 512e DS", "1-Year", "LP MDL"),
        ("P09149-B21", "10TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e DS", "1-Year", "LP MDL"),
        ("834031-B21", "8TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e DS", "1-Year", "LP MDL"),
        ("861746-B21", "6TB", "LFF 3.5\"", "7.2K", "SAS 12G", "512e", "1-Year", "LP MDL"),
    ]
    create_hdd_table_9col(doc, hpe_bc_lff)

    doc.add_page_break()

    # ==================== Contact Page ====================
    contact_bg = doc.add_table(rows=1, cols=1)
    contact_bg.alignment = WD_TABLE_ALIGNMENT.CENTER
    bg_cell = contact_bg.cell(0, 0)
    set_cell_shading(bg_cell, CREAM)
    set_no_border(bg_cell)

    para = bg_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para.add_run("\n\n\n")

    title_run = para.add_run("CONTACT US\n")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(126, 4, 29)
    title_run.font.name = 'Arial'

    line_run = para.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n\n")
    line_run.font.color.rgb = RGBColor(126, 4, 29)

    doc.add_paragraph()

    contact_table = doc.add_table(rows=5, cols=2)
    contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    contact_info = [
        ("Phone", "XXX-XXXX-XXXX"),
        ("WhatsApp", "XXX-XXXX-XXXX"),
        ("Email", "example@company.com"),
        ("Website", "www.yourcompany.com"),
        ("Address", "Your Company Address"),
    ]

    for i, (label, value) in enumerate(contact_info):
        cell0 = contact_table.cell(i, 0)
        set_cell_shading(cell0, BURGUNDY)
        set_no_border(cell0)
        para0 = cell0.paragraphs[0]
        para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = para0.add_run(label)
        run0.font.size = Pt(14)
        run0.font.bold = True
        run0.font.color.rgb = RGBColor(255, 255, 255)
        run0.font.name = 'Arial'

        cell1 = contact_table.cell(i, 1)
        set_cell_shading(cell1, WHITE)
        set_cell_border(cell1, BURGUNDY, "2")
        para1 = cell1.paragraphs[0]
        run1 = para1.add_run("  " + value)
        run1.font.size = Pt(14)
        run1.font.color.rgb = RGBColor(54, 54, 54)
        run1.font.name = 'Arial'

    doc.add_paragraph()

    qr_table = doc.add_table(rows=1, cols=1)
    qr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    qr_cell = qr_table.cell(0, 0)
    set_cell_shading(qr_cell, CREAM)
    set_cell_border(qr_cell, BURGUNDY, "4")
    qr_para = qr_cell.paragraphs[0]
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = qr_para.add_run("\n\n[ QR Code ]\n\n")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(126, 4, 29)
    run.font.name = 'Arial'

    # Save
    output_path = "/home/user/-JC-/Product_Catalog_Burgundy.docx"
    doc.save(output_path)
    print(f"✅ Product Catalog created: {output_path}")
    return output_path

if __name__ == "__main__":
    main()
