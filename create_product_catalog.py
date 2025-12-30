# -*- coding: utf-8 -*-
"""
äº§å“ç›®å½•ç”Ÿæˆè„šæœ¬
åˆ›å»ºç²¾ç¾çš„Wordäº§å“ç›®å½•æ¨¡æ¿ - ä¼˜åŒ–ç‰ˆ
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Babyè“è‰² RGB: 137, 207, 240 -> Hex: 89CFE0
BABY_BLUE = "89CFE0"
BABY_BLUE_LIGHT = "D6EEF8"
BABY_BLUE_DARK = "5DADE2"
DARK_BLUE = "003366"
WHITE = "FFFFFF"

def set_cell_shading(cell, color):
    """è®¾ç½®å•å…ƒæ ¼èƒŒæ™¯é¢œè‰²"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, color="CCCCCC", size="4"):
    """è®¾ç½®å•å…ƒæ ¼è¾¹æ¡†"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_cell_diagonal_border(cell, color="89CFE0"):
    """è®¾ç½®å•å…ƒæ ¼å¯¹è§’çº¿è¾¹æ¡†ï¼ˆåˆ›å»ºè±æ ¼æ•ˆæœï¼‰"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:tl2br w:val="single" w:sz="6" w:color="{color}"/>'
        f'<w:tr2bl w:val="single" w:sz="6" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def create_cover_page(doc):
    """åˆ›å»ºå°é¢é¡µ - ç™½åº•babyè“è±æ ¼è®¾è®¡"""

    # åˆ›å»ºè±æ ¼èƒŒæ™¯è¡¨æ ¼ (8è¡Œx6åˆ—çš„å°æ ¼å­ï¼Œæ¯ä¸ªæ ¼å­æœ‰å¯¹è§’çº¿)
    diamond_table = doc.add_table(rows=8, cols=6)
    diamond_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # è®¾ç½®æ¯ä¸ªå•å…ƒæ ¼çš„å¯¹è§’çº¿è¾¹æ¡†å½¢æˆè±æ ¼
    for row in diamond_table.rows:
        row.height = Cm(1.8)
        for cell in row.cells:
            set_cell_shading(cell, WHITE)
            set_cell_diagonal_border(cell, BABY_BLUE)

    doc.add_paragraph()

    # ä¸»å†…å®¹åŒºåŸŸ - ç™½è‰²èƒŒæ™¯å¡ç‰‡
    content_table = doc.add_table(rows=1, cols=1)
    content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    content_cell = content_table.cell(0, 0)
    set_cell_shading(content_cell, WHITE)
    set_cell_border(content_cell, BABY_BLUE, "12")

    # åœ¨å•å…ƒæ ¼å†…æ·»åŠ å†…å®¹
    para = content_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # æ·»åŠ ç©ºè¡Œ
    para.add_run("\n\n")

    # å…¬å¸Logoå ä½
    logo_run = para.add_run("[ å…¬å¸LOGO ]\n\n")
    logo_run.font.size = Pt(18)
    logo_run.font.color.rgb = RGBColor(137, 207, 240)
    logo_run.font.name = 'Microsoft YaHei'
    logo_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # å…¬å¸åç§°
    company_run = para.add_run("æ‚¨çš„å…¬å¸åç§°\n")
    company_run.font.size = Pt(32)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(0, 51, 102)
    company_run.font.name = 'Microsoft YaHei'
    company_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # åˆ†éš”çº¿æ•ˆæœ
    line_run = para.add_run("â”" * 20 + "\n\n")
    line_run.font.size = Pt(12)
    line_run.font.color.rgb = RGBColor(137, 207, 240)

    # äº§å“ç›®å½•æ ‡é¢˜
    title_run = para.add_run("äº§ å“ ç›® å½•\n")
    title_run.font.size = Pt(42)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(93, 173, 226)
    title_run.font.name = 'Microsoft YaHei'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # è‹±æ–‡æ ‡é¢˜
    eng_run = para.add_run("PRODUCT CATALOG\n\n")
    eng_run.font.size = Pt(16)
    eng_run.font.color.rgb = RGBColor(128, 128, 128)
    eng_run.font.name = 'Arial'

    # å¹´ä»½
    year_run = para.add_run("2024 - 2025\n\n\n")
    year_run.font.size = Pt(14)
    year_run.font.color.rgb = RGBColor(137, 207, 240)

    # è”ç³»ä¿¡æ¯
    contact_run = para.add_run("ç”µè¯ï¼šXXX-XXXX-XXXX | é‚®ç®±ï¼šexample@company.com\n")
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(102, 102, 102)
    contact_run.font.name = 'Microsoft YaHei'
    contact_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    address_run = para.add_run("åœ°å€ï¼šæ‚¨çš„å…¬å¸åœ°å€\n\n")
    address_run.font.size = Pt(10)
    address_run.font.color.rgb = RGBColor(102, 102, 102)
    address_run.font.name = 'Microsoft YaHei'
    address_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # åº•éƒ¨è±æ ¼è£…é¥°
    doc.add_paragraph()
    bottom_table = doc.add_table(rows=3, cols=6)
    bottom_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in bottom_table.rows:
        row.height = Cm(1.2)
        for cell in row.cells:
            set_cell_shading(cell, WHITE)
            set_cell_diagonal_border(cell, BABY_BLUE)

    # åˆ†é¡µ
    doc.add_page_break()

def create_company_intro_page(doc):
    """åˆ›å»ºå…¬å¸ä»‹ç»é¡µ"""
    # æ ‡é¢˜è£…é¥°æ¡
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell0 = header_table.cell(0, 0)
    set_cell_shading(cell0, BABY_BLUE)
    cell0.paragraphs[0].add_run("  ")

    cell1 = header_table.cell(0, 1)
    set_cell_shading(cell1, WHITE)
    para1 = cell1.paragraphs[0]
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = para1.add_run("å…³äºæˆ‘ä»¬ ABOUT US")
    run1.font.size = Pt(24)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0, 51, 102)
    run1.font.name = 'Microsoft YaHei'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    cell2 = header_table.cell(0, 2)
    set_cell_shading(cell2, BABY_BLUE)
    cell2.paragraphs[0].add_run("  ")

    doc.add_paragraph()

    # å…¬å¸ç®€ä»‹æ¡†
    intro_table = doc.add_table(rows=1, cols=1)
    intro_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = intro_table.cell(0, 0)
    set_cell_shading(cell, BABY_BLUE_LIGHT)
    set_cell_border(cell, BABY_BLUE, "8")

    cell_para = cell.paragraphs[0]
    run = cell_para.add_run("""
    å…¬å¸ç®€ä»‹

    [åœ¨æ­¤å¤„å¡«å†™æ‚¨çš„å…¬å¸ç®€ä»‹ï¼ŒåŒ…æ‹¬å…¬å¸æˆç«‹æ—¶é—´ã€ä¸»è¥ä¸šåŠ¡ã€å‘å±•å†ç¨‹ç­‰ã€‚

    å»ºè®®å†…å®¹ï¼š
    â€¢ å…¬å¸æˆç«‹å¹´ä»½åŠå‘å±•å†ç¨‹
    â€¢ ä¸»è¥äº§å“å’ŒæœåŠ¡èŒƒå›´
    â€¢ å…¬å¸è§„æ¨¡å’Œå›¢é˜Ÿå®åŠ›
    â€¢ æ ¸å¿ƒç«äº‰ä¼˜åŠ¿
    â€¢ æœåŠ¡ç†å¿µå’Œç»è¥å®—æ—¨]

    """)
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()
    doc.add_paragraph()

    # æˆ‘ä»¬çš„ä¼˜åŠ¿
    adv_title = doc.add_paragraph()
    adv_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = adv_title.add_run("â—† æˆ‘ä»¬çš„ä¼˜åŠ¿ â—†")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    # ä¼˜åŠ¿è¡¨æ ¼ 2x2
    adv_table = doc.add_table(rows=2, cols=2)
    adv_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    advantages = [
        ("âœ“ æ­£å“ä¿éšœ", "æ‰€æœ‰äº§å“å‡ä¸ºåŸå‚æ­£å“ï¼Œæä¾›å®Œæ•´è´¨ä¿"),
        ("âœ“ ä»·æ ¼ä¼˜åŠ¿", "å‚å®¶ç›´ä¾›ï¼Œä»·æ ¼æ›´å…·ç«äº‰åŠ›"),
        ("âœ“ æŠ€æœ¯æ”¯æŒ", "ä¸“ä¸šæŠ€æœ¯å›¢é˜Ÿï¼Œæä¾›å”®å‰å”®åæœåŠ¡"),
        ("âœ“ å¿«é€Ÿå‘è´§", "å……è¶³åº“å­˜ï¼Œå½“å¤©å‘è´§")
    ]

    for i, (title, desc) in enumerate(advantages):
        row = i // 2
        col = i % 2
        cell = adv_table.cell(row, col)
        set_cell_shading(cell, BABY_BLUE_LIGHT)
        set_cell_border(cell, BABY_BLUE, "6")
        para = cell.paragraphs[0]

        title_run = para.add_run(title + "\n")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title_run.font.name = 'Microsoft YaHei'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        desc_run = para.add_run(desc)
        desc_run.font.size = Pt(11)
        desc_run.font.name = 'Microsoft YaHei'
        desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

def create_packaging_page(doc):
    """åˆ›å»ºåŒ…è£…å±•ç¤ºé¡µ"""
    # æ ‡é¢˜
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("â—† åŒ…è£…å±•ç¤º â—†")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("PACKAGING DISPLAY")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(137, 207, 240)

    doc.add_paragraph()

    # åŒ…è£…å›¾ç‰‡å±•ç¤º - 2è¡Œ3åˆ—å¸ƒå±€
    for row_num in range(2):
        row_table = doc.add_table(rows=2, cols=3)
        row_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # å›¾ç‰‡è¡Œ
        for col in range(3):
            cell = row_table.cell(0, col)
            set_cell_shading(cell, BABY_BLUE_LIGHT)
            set_cell_border(cell, BABY_BLUE, "6")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"\n\n[ åŒ…è£…å›¾ç‰‡ {row_num * 3 + col + 1} ]\n\n")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(93, 173, 226)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # æ ‡é¢˜è¡Œ
        for col in range(3):
            cell = row_table.cell(1, col)
            set_cell_shading(cell, WHITE)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"åŒ…è£…å±•ç¤º {row_num * 3 + col + 1}")
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_paragraph()

    # è¯´æ˜æ–‡å­—
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("æ‰€æœ‰äº§å“å‡é‡‡ç”¨ä¸“ä¸šåŒ…è£…ï¼Œç¡®ä¿è¿è¾“å®‰å…¨")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

def create_catalog_page(doc):
    """åˆ›å»ºäº§å“ç›®å½•ç´¢å¼•é¡µ"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("â—† äº§å“ç›®å½• â—†")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("CONTENTS")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(137, 207, 240)

    doc.add_paragraph()
    doc.add_paragraph()

    # ç›®å½•åˆ—è¡¨
    catalog_items = [
        ("01", "HPE æœåŠ¡å™¨å†…å­˜", "DDR4/DDR5 ç³»åˆ—"),
        ("02", "Dell æœåŠ¡å™¨å†…å­˜", "DDR4/DDR5 ç³»åˆ—"),
        ("03", "Lenovo æœåŠ¡å™¨å†…å­˜", "ThinkSystem ç³»åˆ—"),
        ("04", "HPE ä¼ä¸šçº§ç¡¬ç›˜", "SAS HDD ç³»åˆ—"),
        ("05", "æ›´å¤šäº§å“", "æŒç»­æ›´æ–°ä¸­...")
    ]

    for num, title_text, desc in catalog_items:
        # åˆ›å»ºç›®å½•é¡¹è¡¨æ ¼
        item_table = doc.add_table(rows=1, cols=3)
        item_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # åºå·
        cell0 = item_table.cell(0, 0)
        set_cell_shading(cell0, BABY_BLUE)
        para0 = cell0.paragraphs[0]
        para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = para0.add_run(num)
        run0.font.size = Pt(18)
        run0.font.bold = True
        run0.font.color.rgb = RGBColor(255, 255, 255)

        # æ ‡é¢˜
        cell1 = item_table.cell(0, 1)
        set_cell_shading(cell1, BABY_BLUE_LIGHT)
        para1 = cell1.paragraphs[0]
        run1 = para1.add_run("  " + title_text)
        run1.font.size = Pt(16)
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0, 51, 102)
        run1.font.name = 'Microsoft YaHei'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # æè¿°
        cell2 = item_table.cell(0, 2)
        set_cell_shading(cell2, WHITE)
        set_cell_border(cell2, BABY_BLUE, "4")
        para2 = cell2.paragraphs[0]
        run2 = para2.add_run("  " + desc)
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(102, 102, 102)
        run2.font.name = 'Microsoft YaHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        doc.add_paragraph()

    doc.add_page_break()

def create_product_section_header(doc, title, subtitle, color=BABY_BLUE):
    """åˆ›å»ºäº§å“åˆ†ç±»æ ‡é¢˜"""
    # æ ‡é¢˜èƒŒæ™¯
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_border(cell, color, "0")

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run("  " + title + "  ")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # å‰¯æ ‡é¢˜
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()

def create_product_table(doc, products, columns):
    """åˆ›å»ºäº§å“åˆ—è¡¨è¡¨æ ¼"""
    # è¡¨å¤´
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # è®¾ç½®è¡¨å¤´
    header_row = table.rows[0]
    for i, col_name in enumerate(columns):
        cell = header_row.cells[i]
        set_cell_shading(cell, BABY_BLUE)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(col_name)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # æ·»åŠ äº§å“æ•°æ®
    for idx, product in enumerate(products):
        row = table.add_row()
        bg_color = WHITE if idx % 2 == 0 else BABY_BLUE_LIGHT

        for i, value in enumerate(product):
            cell = row.cells[i]
            set_cell_shading(cell, bg_color)
            set_cell_border(cell, BABY_BLUE, "4")
            para = cell.paragraphs[0]
            run = para.add_run(str(value))
            run.font.size = Pt(9)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

def main():
    """ä¸»å‡½æ•°"""
    doc = Document()

    # è®¾ç½®é¡µé¢è¾¹è·
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # åˆ›å»ºå°é¢
    create_cover_page(doc)

    # åˆ›å»ºå…¬å¸ä»‹ç»é¡µ
    create_company_intro_page(doc)

    # åˆ›å»ºåŒ…è£…å±•ç¤ºé¡µ
    create_packaging_page(doc)

    # åˆ›å»ºäº§å“ç›®å½•ç´¢å¼•
    create_catalog_page(doc)

    # ==================== HPE æœåŠ¡å™¨å†…å­˜ ====================
    create_product_section_header(doc, "HPE æœåŠ¡å™¨å†…å­˜", "HPE Server Memory - DDR4/DDR5 Series")

    # HPE G10 ç³»åˆ—å†…å­˜
    g10_title = doc.add_paragraph()
    run = g10_title.add_run("â–¶ G10 ç³»åˆ—å†…å­˜")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    hpe_g10_products = [
        ("P00924-B21", "HPE 32GB Dual rank x4 DDR4-2933", "32GB", "DDR4-2933", "CAS-21-21-21"),
        ("815100-B21", "HPE 32GB Dual rank x4 DDR4-2666", "32GB", "DDR4-2666", "CAS-19-19-19"),
        ("P00930-B21", "HPE 64GB Dual rank x4 DDR4-2933", "64GB", "DDR4-2933", "CAS-21-21-21"),
        ("P00926-B21", "HPE 64GB Quad rank x4 DDR4-2933", "64GB", "DDR4-2933", "Load reduced"),
        ("P11040-B21", "HPE 128GB Quad rank x4 DDR4-2933", "128GB", "DDR4-2933", "Load reduced"),
        ("815101-B21", "HPE 64GB Quad rank x4 DDR4-2666", "64GB", "DDR4-2666", "Load reduced"),
        ("P00928-B21", "HPE 128GB Octal rank x4 DDR4-2933", "128GB", "DDR4-2933", "3DS"),
        ("815102-B21", "HPE 128GB Octal rank x4 DDR4-2666", "128GB", "DDR4-2666", "3DS"),
    ]

    create_product_table(doc, hpe_g10_products, ["å‹å·", "æè¿°", "å®¹é‡", "è§„æ ¼", "ç±»å‹"])

    # HPE G10+ ç³»åˆ—å†…å­˜
    g10plus_title = doc.add_paragraph()
    run = g10plus_title.add_run("â–¶ G10+ ç³»åˆ—å†…å­˜")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    hpe_g10plus_products = [
        ("P06033-B21", "HPE 32GB Dual rank x4 DDR4-3200", "32GB", "DDR4-3200", "CAS-22-22-22"),
        ("P06035-B21", "HPE 64GB Dual rank x4 DDR4-3200", "64GB", "DDR4-3200", "CAS-22-22-22"),
        ("P40007-B21", "HPE 32GB Single rank x4 DDR4-3200", "32GB", "DDR4-3200", "Registered"),
        ("P06037-B21", "HPE 128GB Quad rank x4 DDR4-3200", "128GB", "DDR4-3200", "Load reduced"),
        ("P06039-B21", "HPE 256GB Octal rank x4 DDR4-3200", "256GB", "DDR4-3200", "3DS"),
    ]

    create_product_table(doc, hpe_g10plus_products, ["å‹å·", "æè¿°", "å®¹é‡", "è§„æ ¼", "ç±»å‹"])

    # HPE G11 ç³»åˆ—å†…å­˜ (DDR5)
    g11_title = doc.add_paragraph()
    run = g11_title.add_run("â–¶ G11 ç³»åˆ—å†…å­˜ (DDR5)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    hpe_g11_products = [
        ("P64705-B21", "HPE 16GB Single Rank x8 DDR5-5600", "16GB", "DDR5-5600", "P65253-001"),
        ("P64706-B21", "HPE 32GB Dual Rank x8 DDR5-5600", "32GB", "DDR5-5600", "P65254-001"),
        ("P64707-B21", "HPE 64GB Dual Rank x4 DDR5-5600", "64GB", "DDR5-5600", "P65255-001"),
        ("P64708-B21", "HPE 96GB Dual Rank x4 DDR5-5600", "96GB", "DDR5-5600", "P65256-001"),
        ("P64709-B21", "HPE 128GB Quad Rank x4 DDR5-5600", "128GB", "DDR5-5600", "P65257-001"),
        ("P43322-B21", "HPE 16GB Single Rank x8 DDR5-4800", "16GB", "DDR5-4800", "P48499-001"),
        ("P43328-B21", "HPE 32GB Dual Rank x8 DDR5-4800", "32GB", "DDR5-4800", "P48501-001"),
        ("P43331-B21", "HPE 64GB Dual Rank x4 DDR5-4800", "64GB", "DDR5-4800", "P48502-001"),
        ("P66675-B21", "HPE 96GB 2Rx4 PC5-4800B-R Smart Kit", "96GB", "DDR5-4800", "P67364-001"),
        ("P43334-B21", "HPE 128GB Quad Rank x4 DDR5-4800", "128GB", "DDR5-4800", "P48503-001"),
        ("P43337-B21", "HPE 256GB Octal Rank x4 DDR5-4800", "256GB", "DDR5-4800", "P48504-001"),
    ]

    create_product_table(doc, hpe_g11_products, ["å‹å·", "æè¿°", "å®¹é‡", "è§„æ ¼", "å¤‡ä»¶å·"])

    doc.add_page_break()

    # ==================== Dell æœåŠ¡å™¨å†…å­˜ ====================
    create_product_section_header(doc, "Dell æœåŠ¡å™¨å†…å­˜", "Dell Server Memory Series", "5DADE2")

    dell_products = [
        ("AA601616", "SNP8WKDYC/32G 32GB PC4-23400 DDR4-2933MHz", "32GB", "DDR4-2933", "2Rx4"),
        ("AA601615", "SNPW403YC/64G 64GB DDR4 RDIMM 2933 MT/s", "64GB", "DDR4-2933", "2Rx4"),
        ("AA579531", "SNP8WKDYC/32G 32-GB 2Rx4 2933MHz PC4-23400", "32GB", "DDR4-2933", "2Rx4"),
        ("AA579530", "SNPW403YC/64G 64GB DDR4 RDIMM 2933 MT/s", "64GB", "DDR4-2933", "2Rx4"),
        ("AA783422", "SNP75X1VC/32G 32GB DDR4 RDIMM 3200 MT/s", "32GB", "DDR4-3200", "2Rx4"),
        ("AA783423", "SNPP2MYX/64G 64GB DDR4 RDIMM 3200 MT/s", "64GB", "DDR4-3200", "2Rx4"),
        ("AA799110", "SNPP2MYX/64G 64GB DDR4 RDIMM 3200 MT/s", "64GB", "DDR4-3200", "2Rx4"),
        ("AA810828", "SNPP2MYX/64G 64GB DDR4 RDIMM 3200 MT/s", "64GB", "DDR4-3200", "2Rx4"),
        ("AB445285", "SNP7JXF5C/128G 128GB DDR4 LRDIMM 3200 MT/s", "128GB", "DDR4-3200", "4Rx4"),
        ("AC239377", "SNP1V1N1C/16G 16GB DDR5 RDIMM 4800 MT/s", "16GB", "DDR5-4800", "1Rx8"),
        ("AC239378", "SNPW08W9C/32G 32GB DDR5 RDIMM 4800 MT/s", "32GB", "DDR5-4800", "2Rx8"),
        ("AC239379", "SNP152K5C/64G 64GB DDR5 RDIMM 4800 MT/s", "64GB", "DDR5-4800", "2Rx4"),
        ("AC830716", "SNPSD48RC/16G 16GB PC5-44800 DDR5-5600MHz", "16GB", "DDR5-5600", "1Rx8"),
        ("AC958788", "SNPXHJ68MC/16G AC958788 16G DDR5 UDIMM 5600 MT/s", "16GB", "DDR5-5600", "1Rx8"),
        ("AC774043", "SNP8D9M0C/32G 32G 2RX8 PC5 5600B-E DDR5 ECC UDIMM", "32GB", "DDR5-5600", "2Rx8"),
        ("AC830717", "SNPP8XPWC/32G 32GB 2RX8 PC5-44800 DDR5-5600B ECC RDIMM", "32GB", "DDR5-5600", "2Rx8"),
        ("AC830718", "SNP58F3NC/64G 64GB 2Rx4 PC5-44800B-R DDR5-5600 RDIMM", "64GB", "DDR5-5600", "2Rx4"),
        ("AC888060", "SNP5DR48C/16G 16GB DDR5-5600 ECC RDIMM", "16GB", "DDR5-5600", "1Rx8"),
    ]

    create_product_table(doc, dell_products, ["å‹å·", "æè¿°", "å®¹é‡", "è§„æ ¼", "Rank"])

    doc.add_page_break()

    # ==================== Lenovo æœåŠ¡å™¨å†…å­˜ ====================
    create_product_section_header(doc, "Lenovo æœåŠ¡å™¨å†…å­˜", "Lenovo ThinkSystem Memory Series", "3498DB")

    lenovo_products = [
        ("4ZC7A08707", "ThinkSystem 01KR353 16GB 1Rx4 PC4-2933Y", "16GB", "DDR4-2933", "1Rx4"),
        ("4ZC7A08709", "ThinkSystem 01KR355 32GB 2RX4 PC4-2933Y DDR4 REG ECC", "32GB", "DDR4-2933", "2Rx4"),
        ("4ZC7A08710", "ThinkSystem 01KR356 64GB 2Rx4 PC4-2933Y RECC", "64GB", "DDR4-2933", "2Rx4"),
        ("4X77A08634", "ThinkSystem 02JK239 32GB 2RX8 DDR4 3200 RDIMM", "32GB", "DDR4-3200", "2Rx8"),
        ("4X77A08633", "ThinkSystem 02JK237 32GB 2Rs4 PC4-3200AA", "32GB", "DDR4-3200", "2Rs4"),
        ("4ZC7A15124", "ThinkSystem 02JG340 64GB 2RX4 PC4-3200AA-RDIMM", "64GB", "DDR4-3200", "2Rx4"),
        ("4X77A08635", "ThinkSystem 02JK971 64G 2RX4 DDR4 3200", "64GB", "DDR4-3200", "2Rx4"),
        ("4X77A77496", "ThinkSystem 03GX401 32GB PC4-3200AA ECC UDIMM", "32GB", "DDR4-3200", "ECC UDIMM"),
        ("4X77A85511", "ThinkSystem 16GB TruDDR5 4800 MHz (1Rx8)", "16GB", "DDR5-4800", "1Rx8"),
        ("4X77A77483", "ThinkSystem 32GB TruDDR5 4800MHz (1Rx4)", "32GB", "DDR5-4800", "1Rx4"),
        ("4X77A88512", "ThinkSystem 32GB TruDDR5 4800MHz (2Rx8)", "32GB", "DDR5-4800", "2Rx8"),
        ("4X77A77031", "ThinkSystem 32GB TruDDR5 4800MHz (2Rx8)", "32GB", "DDR5-4800", "2Rx8"),
        ("4X77A81440", "ThinkSystem 03KL461 32G TruDDR5 2RX8 4800", "32GB", "DDR5-4800", "2Rx8"),
        ("4X77A81442", "ThinkSystem 03GX338 64G TruDDR5 2RX4 PC5 4800B", "64GB", "DDR5-4800", "2Rx4"),
        ("4X77A77033", "ThinkSystem 64GB TruDDR5 4800MHz (2Rx4)", "64GB", "DDR5-4800", "2Rx4"),
        ("4X77A77034", "ThinkSystem 128GB TruDDR5 4800MHz 4Rx4 3DS RDIMM", "128GB", "DDR5-4800", "4Rx4"),
        ("4X77A77035", "ThinkSystem 256GB TruDDR5 4800MHz (8Rx4)", "256GB", "DDR5-4800", "8Rx4"),
        ("4X77A88049", "ThinkSystem 32GB TruDDR5 5600MHz (1Rx4)", "32GB", "DDR5-5600", "1Rx4"),
        ("4X77A88051", "ThinkSystem 32GB TruDDR5 5600MHz (2Rx8)", "32GB", "DDR5-5600", "2Rx8"),
        ("4X77A90992", "ThinkSystem 64GB TruDDR5 5600MHz 2Rx4", "64GB", "DDR5-5600", "2Rx4"),
        ("4X77A88052", "ThinkSystem 64GB TruDDR5 5600MHz (2Rx4)", "64GB", "DDR5-5600", "2Rx4"),
        ("4X77A93887", "ThinkSystem 128GB TruDDR5 5600MHz (2Rx4) RDIMM", "128GB", "DDR5-5600", "2Rx4"),
        ("4X77A88054", "ThinkSystem 128GB TruDDR5 5600MHz (4Rx4)", "128GB", "DDR5-5600", "4Rx4"),
        ("4X77A88055", "ThinkSystem 256GB TruDDR5 5600 MHz (8Rx4) 3DS RDIMM", "256GB", "DDR5-5600", "8Rx4"),
    ]

    create_product_table(doc, lenovo_products, ["å‹å·", "æè¿°", "å®¹é‡", "è§„æ ¼", "Rank"])

    doc.add_page_break()

    # ==================== HPE ä¼ä¸šçº§ç¡¬ç›˜ ====================
    create_product_section_header(doc, "HPE ä¼ä¸šçº§ç¡¬ç›˜", "HPE Enterprise SAS HDD Series", "2E86C1")

    # ä¼ä¸šçº§ 10K/15K ç¡¬ç›˜
    enterprise_title = doc.add_paragraph()
    run = enterprise_title.add_run("â–¶ ä¼ä¸šçº§ 10K/15K SAS ç¡¬ç›˜ (2.5å¯¸ SFF)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    hpe_hdd_enterprise = [
        ("881457-B21", "HPE 2.4TB SAS 12G Enterprise 10K SFF 512e DS firmware HDD", "2.4TB", "10K", "SFF 2.5å¯¸"),
        ("872481-B21", "HPE 1.8TB SAS 12G Enterprise 10K SFF 512e DS firmware HDD", "1.8TB", "10K", "SFF 2.5å¯¸"),
        ("872479-B21", "HPE 1.2TB SAS 12G Enterprise 10K SFF DS firmware HDD", "1.2TB", "10K", "SFF 2.5å¯¸"),
        ("870759-B21", "HPE 900GB SAS 12G Enterprise 15K SFF DS firmware HDD", "900GB", "15K", "SFF 2.5å¯¸"),
        ("870757-B21", "HPE 600GB SAS 12G Enterprise 15K SFF DS firmware HDD", "600GB", "15K", "SFF 2.5å¯¸"),
        ("872477-B21", "HPE 600GB SAS 12G Enterprise 10K SFF DS firmware HDD", "600GB", "10K", "SFF 2.5å¯¸"),
        ("870753-B21", "HPE 300GB SAS 12G Enterprise 15K SFF DS firmware HDD", "300GB", "15K", "SFF 2.5å¯¸"),
        ("872475-B21", "HPE 300GB SAS 12G Enterprise 10K SFF DS firmware HDD", "300GB", "10K", "SFF 2.5å¯¸"),
    ]

    create_product_table(doc, hpe_hdd_enterprise, ["å‹å·", "æè¿°", "å®¹é‡", "è½¬é€Ÿ", "å°ºå¯¸"])

    # Midline 7.2K ç¡¬ç›˜
    midline_title = doc.add_paragraph()
    run = midline_title.add_run("â–¶ Midline 7.2K SAS ç¡¬ç›˜")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    hpe_hdd_midline = [
        ("765466-B21", "HPE 2TB SAS 12G Midline 7.2K SFF 512e HDD", "2TB", "7.2K", "SFF 2.5å¯¸"),
        ("832514-B21", "HPE 1TB SAS 12G Midline 7.2K SFF DS firmware HDD", "1TB", "7.2K", "SFF 2.5å¯¸"),
        ("P23863-B21", "HPE 16TB SAS 12G Business Critical 7.2K LFF 512e ISE HDD", "16TB", "7.2K", "LFF 3.5å¯¸"),
        ("P09153-B21", "HPE 14TB SAS 12G Midline 7.2K LFF Helium 512e DS firmware HDD", "14TB", "7.2K", "LFF 3.5å¯¸"),
        ("881779-B21", "HPE 12TB SAS 12G Midline 7.2K LFF Helium 512e DS firmware HDD", "12TB", "7.2K", "LFF 3.5å¯¸"),
        ("857644-B21", "HPE 10TB SAS 12G Midline 7.2K LFF Helium 512e DS firmware HDD", "10TB", "7.2K", "LFF 3.5å¯¸"),
        ("819201-B21", "HPE 8TB SAS 12G Midline 7.2K LFF DS firmware HDD", "8TB", "7.2K", "LFF 3.5å¯¸"),
        ("861754-B21", "HPE 6TB SAS 12G Midline 7.2K LFF HDD", "6TB", "7.2K", "LFF 3.5å¯¸"),
    ]

    create_product_table(doc, hpe_hdd_midline, ["å‹å·", "æè¿°", "å®¹é‡", "è½¬é€Ÿ", "å°ºå¯¸"])

    # Mission Critical ç¡¬ç›˜
    mc_title = doc.add_paragraph()
    run = mc_title.add_run("â–¶ Mission Critical SAS ç¡¬ç›˜")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    hpe_hdd_mc = [
        ("P28352-B21", "HPE 2.4TB SAS 12G mission critical 10K SFF BC 3-year warranty 512e HDD", "2.4TB", "10K", "SFF"),
        ("P28586-B21", "HPE 1.2TB SAS 12G mission critical 10K SFF BC 3-year warranty HDD", "1.2TB", "10K", "SFF"),
        ("P28028-B21", "HPE 300GB SAS 12G mission critical 15K SFF BC 3-year warranty HDD", "300GB", "15K", "SFF"),
        ("P40430-B21", "HPE 300GB SAS 12G mission critical 10K SFF BC 3-year warranty HDD", "300GB", "10K", "SFF"),
        ("P37669-B21", "HPE 18TB SAS 12G business critical 7.2K LFF LP 1-year warranty 512e ISE HDD", "18TB", "7.2K", "LFF"),
        ("P23608-B21", "HPE 16TB SAS 12G business critical 7.2K LFF LP 1-year warranty 512e ISE HDD", "16TB", "7.2K", "LFF"),
        ("P09155-B21", "HPE 14TB SAS 12G MDL 7.2K LFF LP 1-year warranty helium 512e DS firmware HDD", "14TB", "7.2K", "LFF"),
        ("881781-B21", "HPE 12TB SAS 12G MDL 7.2K LFF LP 1-year warranty helium 512e DS firmware HDD", "12TB", "7.2K", "LFF"),
        ("P09149-B21", "HPE 10TB SAS 12G MDL 7.2K LFF LP 1-year warranty helium 512e DS firmware HDD", "10TB", "7.2K", "LFF"),
        ("834031-B21", "HPE 8TB SAS 12G MDL 7.2K LFF LP 1-year warranty 512e DS firmware HDD", "8TB", "7.2K", "LFF"),
        ("861746-B21", "HPE 6TB SAS 12G MDL 7.2K LFF LP 1-year warranty 512e HDD", "6TB", "7.2K", "LFF"),
    ]

    create_product_table(doc, hpe_hdd_mc, ["å‹å·", "æè¿°", "å®¹é‡", "è½¬é€Ÿ", "å°ºå¯¸"])

    doc.add_page_break()

    # ==================== è”ç³»æˆ‘ä»¬é¡µé¢ ====================
    # é¡¶éƒ¨è£…é¥°
    for _ in range(2):
        doc.add_paragraph()

    top_decor = doc.add_table(rows=2, cols=6)
    top_decor.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in top_decor.rows:
        row.height = Cm(1)
        for cell in row.cells:
            set_cell_shading(cell, WHITE)
            set_cell_diagonal_border(cell, BABY_BLUE)

    contact_title = doc.add_paragraph()
    contact_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_title.add_run("â—† è”ç³»æˆ‘ä»¬ â—†")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    contact_sub = doc.add_paragraph()
    contact_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_sub.add_run("CONTACT US")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(137, 207, 240)

    doc.add_paragraph()

    # è”ç³»ä¿¡æ¯è¡¨æ ¼
    contact_table = doc.add_table(rows=5, cols=2)
    contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    contact_info = [
        ("ğŸ“ ç”µè¯", "XXX-XXXX-XXXX"),
        ("ğŸ“± æ‰‹æœº", "XXX-XXXX-XXXX"),
        ("âœ‰ï¸ é‚®ç®±", "example@company.com"),
        ("ğŸŒ ç½‘å€", "www.yourcompany.com"),
        ("ğŸ“ åœ°å€", "æ‚¨çš„å…¬å¸è¯¦ç»†åœ°å€"),
    ]

    for i, (label, value) in enumerate(contact_info):
        cell0 = contact_table.cell(i, 0)
        set_cell_shading(cell0, BABY_BLUE_LIGHT)
        set_cell_border(cell0, BABY_BLUE, "4")
        para0 = cell0.paragraphs[0]
        run0 = para0.add_run(label)
        run0.font.size = Pt(14)
        run0.font.bold = True
        run0.font.color.rgb = RGBColor(0, 51, 102)
        run0.font.name = 'Microsoft YaHei'
        run0._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        cell1 = contact_table.cell(i, 1)
        set_cell_shading(cell1, WHITE)
        set_cell_border(cell1, BABY_BLUE, "4")
        para1 = cell1.paragraphs[0]
        run1 = para1.add_run(value)
        run1.font.size = Pt(14)
        run1.font.name = 'Microsoft YaHei'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    # äºŒç»´ç å ä½
    qr_table = doc.add_table(rows=1, cols=1)
    qr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    qr_cell = qr_table.cell(0, 0)
    set_cell_shading(qr_cell, BABY_BLUE_LIGHT)
    set_cell_border(qr_cell, BABY_BLUE, "6")
    qr_para = qr_cell.paragraphs[0]
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = qr_para.add_run("\n\n[ å¾®ä¿¡äºŒç»´ç  ]\n\n")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(93, 173, 226)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    scan_para = doc.add_paragraph()
    scan_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = scan_para.add_run("æ‰«ç æ·»åŠ å¾®ä¿¡")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # åº•éƒ¨è£…é¥°
    doc.add_paragraph()
    bottom_decor = doc.add_table(rows=2, cols=6)
    bottom_decor.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in bottom_decor.rows:
        row.height = Cm(1)
        for cell in row.cells:
            set_cell_shading(cell, WHITE)
            set_cell_diagonal_border(cell, BABY_BLUE)

    # ä¿å­˜æ–‡æ¡£
    output_path = "/home/user/-JC-/äº§å“ç›®å½•_Product_Catalog.docx"
    doc.save(output_path)
    print(f"âœ… äº§å“ç›®å½•å·²æˆåŠŸåˆ›å»º: {output_path}")
    return output_path

if __name__ == "__main__":
    main()
