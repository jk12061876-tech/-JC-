# -*- coding: utf-8 -*-
"""
Business Card Generator
Premium Color Scheme - Cream/Burgundy/Gray
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Color Scheme
CREAM = "EFDFCE"          # Cream/Beige - Background
BURGUNDY = "7E041D"       # Burgundy Red - Accent
DARK_GRAY = "363636"      # Dark Gray - Text
WHITE = "FFFFFF"

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="nil"/>'
        f'<w:left w:val="nil"/>'
        f'<w:bottom w:val="nil"/>'
        f'<w:right w:val="nil"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def create_business_card(doc):
    """Create a professional business card"""

    # Main card container
    card = doc.add_table(rows=1, cols=2)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left side - Burgundy accent bar
    left_cell = card.cell(0, 0)
    set_cell_shading(left_cell, BURGUNDY)
    set_no_border(left_cell)
    left_cell.width = Cm(1.5)

    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_para.add_run("\n\n\n\n\n\n\n\n")

    # Right side - Main content on cream background
    right_cell = card.cell(0, 1)
    set_cell_shading(right_cell, CREAM)
    set_no_border(right_cell)

    para = right_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Logo placeholder
    para.add_run("\n\n")
    logo = para.add_run("  [ COMPANY LOGO ]\n\n")
    logo.font.size = Pt(12)
    logo.font.color.rgb = RGBColor(126, 4, 29)
    logo.font.name = 'Arial'

    # Company Name
    company = para.add_run("  YOUR COMPANY NAME\n")
    company.font.size = Pt(18)
    company.font.bold = True
    company.font.color.rgb = RGBColor(54, 54, 54)
    company.font.name = 'Arial'

    # Tagline
    tagline = para.add_run("  Server Components & IT Solutions\n\n")
    tagline.font.size = Pt(10)
    tagline.font.italic = True
    tagline.font.color.rgb = RGBColor(126, 4, 29)
    tagline.font.name = 'Arial'

    # Divider
    divider = para.add_run("  ─────────────────────────────\n\n")
    divider.font.size = Pt(8)
    divider.font.color.rgb = RGBColor(126, 4, 29)

    # Name
    name = para.add_run("  John Smith\n")
    name.font.size = Pt(16)
    name.font.bold = True
    name.font.color.rgb = RGBColor(54, 54, 54)
    name.font.name = 'Arial'

    # Title
    title = para.add_run("  Sales Manager\n\n")
    title.font.size = Pt(11)
    title.font.color.rgb = RGBColor(126, 4, 29)
    title.font.name = 'Arial'

    # Contact Info
    contacts = [
        ("📞  ", "+86 XXX-XXXX-XXXX"),
        ("📱  ", "+86 XXX-XXXX-XXXX (WhatsApp)"),
        ("✉️  ", "email@yourcompany.com"),
        ("🌐  ", "www.yourcompany.com"),
        ("📍  ", "Your Company Address"),
    ]

    for icon, info in contacts:
        line = para.add_run(f"  {icon}{info}\n")
        line.font.size = Pt(10)
        line.font.color.rgb = RGBColor(54, 54, 54)
        line.font.name = 'Arial'

    para.add_run("\n")

    doc.add_paragraph()
    doc.add_paragraph()

    # ============ Card Style 2 - Horizontal Layout ============

    style2_title = doc.add_paragraph()
    style2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = style2_title.add_run("─── Alternative Design ───")
    t.font.size = Pt(12)
    t.font.color.rgb = RGBColor(126, 4, 29)

    doc.add_paragraph()

    # Card 2 - Top bar design
    card2 = doc.add_table(rows=2, cols=1)
    card2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Top burgundy bar
    top_cell = card2.cell(0, 0)
    set_cell_shading(top_cell, BURGUNDY)
    set_no_border(top_cell)

    top_para = top_cell.paragraphs[0]
    top_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    top_para.add_run("\n")
    company2 = top_para.add_run("YOUR COMPANY NAME")
    company2.font.size = Pt(20)
    company2.font.bold = True
    company2.font.color.rgb = RGBColor(255, 255, 255)
    company2.font.name = 'Arial'

    tagline2 = top_para.add_run("\nServer Components & IT Solutions\n")
    tagline2.font.size = Pt(10)
    tagline2.font.color.rgb = RGBColor(239, 223, 206)
    tagline2.font.name = 'Arial'

    # Bottom cream area
    bottom_cell = card2.cell(1, 0)
    set_cell_shading(bottom_cell, CREAM)
    set_no_border(bottom_cell)

    bottom_para = bottom_cell.paragraphs[0]
    bottom_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bottom_para.add_run("\n\n")

    name2 = bottom_para.add_run("John Smith\n")
    name2.font.size = Pt(18)
    name2.font.bold = True
    name2.font.color.rgb = RGBColor(54, 54, 54)
    name2.font.name = 'Arial'

    title2 = bottom_para.add_run("Sales Manager\n\n")
    title2.font.size = Pt(12)
    title2.font.color.rgb = RGBColor(126, 4, 29)
    title2.font.name = 'Arial'

    # Contact row
    contact2 = bottom_para.add_run("📞 +86 XXX-XXXX-XXXX    📱 WhatsApp: +86 XXX-XXXX-XXXX\n")
    contact2.font.size = Pt(10)
    contact2.font.color.rgb = RGBColor(54, 54, 54)
    contact2.font.name = 'Arial'

    email2 = bottom_para.add_run("✉️ email@yourcompany.com    🌐 www.yourcompany.com\n")
    email2.font.size = Pt(10)
    email2.font.color.rgb = RGBColor(54, 54, 54)
    email2.font.name = 'Arial'

    addr2 = bottom_para.add_run("📍 Your Company Address\n\n")
    addr2.font.size = Pt(10)
    addr2.font.color.rgb = RGBColor(54, 54, 54)
    addr2.font.name = 'Arial'

    doc.add_paragraph()
    doc.add_paragraph()

    # ============ Card Style 3 - Minimal ============

    style3_title = doc.add_paragraph()
    style3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3 = style3_title.add_run("─── Minimal Design ───")
    t3.font.size = Pt(12)
    t3.font.color.rgb = RGBColor(126, 4, 29)

    doc.add_paragraph()

    card3 = doc.add_table(rows=1, cols=1)
    card3.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell3 = card3.cell(0, 0)
    set_cell_shading(cell3, CREAM)
    set_no_border(cell3)

    p3 = cell3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3.add_run("\n\n")

    # Burgundy dot accent
    dot = p3.add_run("●\n")
    dot.font.size = Pt(24)
    dot.font.color.rgb = RGBColor(126, 4, 29)

    name3 = p3.add_run("JOHN SMITH\n")
    name3.font.size = Pt(22)
    name3.font.bold = True
    name3.font.color.rgb = RGBColor(54, 54, 54)
    name3.font.name = 'Arial'

    title3 = p3.add_run("Sales Manager\n")
    title3.font.size = Pt(12)
    title3.font.color.rgb = RGBColor(126, 4, 29)
    title3.font.name = 'Arial'

    company3 = p3.add_run("Your Company Name\n\n")
    company3.font.size = Pt(11)
    company3.font.color.rgb = RGBColor(54, 54, 54)
    company3.font.name = 'Arial'

    line3 = p3.add_run("━━━━━━━━━━━━━━━━━━\n\n")
    line3.font.size = Pt(8)
    line3.font.color.rgb = RGBColor(126, 4, 29)

    info3 = p3.add_run("+86 XXX-XXXX-XXXX  |  email@company.com\n")
    info3.font.size = Pt(10)
    info3.font.color.rgb = RGBColor(54, 54, 54)
    info3.font.name = 'Arial'

    web3 = p3.add_run("www.yourcompany.com\n\n")
    web3.font.size = Pt(10)
    web3.font.color.rgb = RGBColor(126, 4, 29)
    web3.font.name = 'Arial'

def main():
    doc = Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("BUSINESS CARD TEMPLATES")
    t.font.size = Pt(24)
    t.font.bold = True
    t.font.color.rgb = RGBColor(126, 4, 29)
    t.font.name = 'Arial'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Choose your favorite style and customize all content")
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(54, 54, 54)
    s.font.name = 'Arial'

    doc.add_paragraph()

    # Create cards
    create_business_card(doc)

    # Save
    output_path = "/home/user/-JC-/Business_Card.docx"
    doc.save(output_path)
    print(f"✅ Business Card created: {output_path}")
    return output_path

if __name__ == "__main__":
    main()
